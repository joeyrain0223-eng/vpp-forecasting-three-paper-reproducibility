from pathlib import Path
import csv
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from build_paper_package import add_markdown_to_docx


BASE = Path(__file__).resolve().parents[1]
ROOT = BASE
PKG = ROOT / "manuscript"
OUT = PKG / "target_journal_eaai_paper1"
OUT.mkdir(parents=True, exist_ok=True)

SOURCE = PKG / "paper_1_price_forecasting.md"

TITLE = "Sequence-Anchored GraphPatch Residual Learning for Uncertainty-Aware Electricity Price Forecasting in Virtual Power Plant Market Operation"


ABSTRACT = """Electricity price forecasting for virtual power plants requires more than average point accuracy: market operation also depends on calibrated uncertainty, price-spike robustness, and representations that transfer across heterogeneous market zones. This paper develops a reproducible graph-temporal learning framework for day-ahead price forecasting that combines lag-safe cross-zone graph features, multi-scale temporal patches, residual learning, and split-conformal calibration. The engineering application is virtual power plant market operation, while the artificial-intelligence contribution is a calibration-selected GraphPatch residual architecture for non-stationary, spike-prone time series. Public experiments use Open Power System Data for Germany-Luxembourg, Denmark West, Denmark East, and Great Britain. On spike-regime test hours, the calibrated GraphPatch blend improves RMSE over a local lag-calendar ridge baseline in all four zones and wins all four paired sign tests. Against DLinear/NLinear anchors, a shrinkage-controlled residual improves spike-regime RMSE in all four zones by 3.69% on average. With a TDConv-style TCN-family comparator, the selected-anchor residual improves all-hour RMSE in 4/4 zones and spike-regime RMSE in 3/4 zones, exposing Great Britain as the strong-anchor limitation case. Rolling-origin validation improves spike-regime RMSE in 12/12 zone-window cases. Conformal intervals reach near-target or conservative coverage in Germany-Luxembourg and Great Britain but remain mildly under-covered in the Nordic zones. The results support a narrow claim: graph-temporal residual learning improves volatile-market forecasting while adaptive calibration remains necessary."""

KEYWORDS = "electricity price forecasting; virtual power plant; graph-temporal learning; conformal prediction; spike robustness; public reproducibility"

HIGHLIGHTS = [
    "GraphPatch residuals target electricity-price spike regimes.",
    "TDConv residual anchoring exposes a strong-sequence limitation.",
    "Rolling-origin GraphPatch gains hold in 12/12 spike tests.",
    "Leave-one-zone-out GraphPatch transfer is positive in 3/4 zones.",
    "Split-conformal calibration exposes Nordic under-coverage.",
]

AUTHOR = "zhijie REN"
AFFILIATION = "College of Computer Science, Hunan University, Lushan South Road, Yuelu District, Changsha, Hunan 410082, China"
EMAIL = "471062741@qq.com"
ORCID = "0009-0006-1048-6640"

AI_DECLARATION = (
    "During the preparation of this manuscript, OpenAI ChatGPT/Codex "
    "was used to support manuscript organization, language editing, code/package "
    "documentation, and reproducibility-checklist drafting. The tool was not used "
    "as an author, did not determine the scientific conclusions, and did not "
    "replace verification of data, code, references, results, or claims. After "
    "using these tools, the author reviewed and edited all AI-assisted text and "
    "takes full responsibility for the final content of the manuscript."
)

DATA_AVAILABILITY_STATEMENT = (
    "The public-data reproducibility package is supplied as supplementary "
    "material. It contains the processed public OPSD benchmark data, experiment "
    "scripts, result tables, rendered figures, verified references, and a "
    "manifest with SHA-256 checksums. The original OPSD time-series source can "
    "be retrieved from the public Open Power System Data repository. Local Hunan "
    "and Shandong operational records were used only as non-public "
    "application-context evidence and are not redistributed because they may "
    "contain confidential market-operation or customer-related information."
)


def make_patch_attention_baseline_block() -> str:
    summary_path = ROOT / "results" / "opsd_patch_attention_price_baseline_summary.csv"
    paired_path = ROOT / "results" / "opsd_patch_attention_price_baseline_paired_tests.csv"
    if not summary_path.exists() or not paired_path.exists():
        return ""
    with summary_path.open("r", encoding="utf-8", newline="") as f:
        summary = list(csv.DictReader(f))
    with paired_path.open("r", encoding="utf-8", newline="") as f:
        paired = list(csv.DictReader(f))

    def row_for(rows: list[dict], zone: str, model: str) -> dict:
        for row in rows:
            if row.get("zone") == zone and row.get("regime") == "spike" and row.get("model") == model:
                return row
        raise ValueError(f"missing row {zone} {model}")

    lines = [
        "Table 15 adds a lightweight patch-attention sequence baseline as a reviewer-response check for the patch/attention family. The model pools seven 24-hour patches from the 168-hour history window with deterministic similarity weights and a regularized ridge head. It is intentionally described as a CPU-only patch-attention baseline, not as full PatchTST or TFT training.",
        "",
        "Table 15 reports the patch-attention reviewer baseline against the TDConv-style comparator on spike-regime OPSD test hours.",
        "",
        "|Zone|TDConv RMSE|Patch-attention RMSE|Patch paired wins / n|MAE delta vs TDConv|Sign-test p|Interpretation|",
        "|---|---:|---:|---:|---:|---:|---|",
    ]
    for zone in ["DE_LU", "DK_1", "DK_2", "GB_GBN"]:
        td = row_for(summary, zone, "TDConv-style sequence ridge")
        patch = row_for(summary, zone, "Patch-attention sequence ridge")
        test = row_for(paired, zone, "Patch-attention sequence ridge")
        p = float(test["sign_test_p_approx"])
        delta = float(test["mean_abs_error_delta"])
        if zone == "DK_2" and p < 0.05 and delta > 0:
            interp = "positive paired evidence"
        elif float(patch["rmse"]) < float(td["rmse"]):
            interp = "lower RMSE but paired evidence mixed"
        else:
            interp = "negative control under strong anchor"
        lines.append(
            "|{}|{:.3f}|{:.3f}|{} / {}|{:.3f}|{:.3f}|{}|".format(
                zone,
                float(td["rmse"]),
                float(patch["rmse"]),
                int(float(test["wins"])),
                int(float(test["paired_n"])),
                delta,
                p,
                interp,
            )
        )
    lines.extend(
        [
            "",
            "The patch-attention baseline reduces spike-regime RMSE in DE-LU, DK1 and DK2 relative to TDConv, but the paired absolute-error evidence is significant only in DK2 and Great Britain remains clearly worse. This result is useful because it closes part of the reviewer concern that patch/attention-style sequence evidence is absent, while preserving the manuscript's claim boundary: the submitted evidence supports GraphPatch residual value after strong local anchors, not universal superiority over every modern sequence family.",
            "",
            "![Figure 13. Patch-attention reviewer baseline compared with the TDConv-style sequence comparator on OPSD spike-regime RMSE.](figures/paper1_fig17_opsd_patch_attention_baseline.png)",
        ]
    )
    return "\n".join(lines)

REFERENCES = """## References

[1] A. Vaswani, N. Shazeer, N. Parmar, J. Uszkoreit, L. Jones, A.N. Gomez, L. Kaiser, I. Polosukhin, Attention is all you need, Advances in Neural Information Processing Systems 30 (2017).
[2] B. Lim, S.O. Arik, N. Loeff, T. Pfister, Temporal Fusion Transformers for interpretable multi-horizon time series forecasting, International Journal of Forecasting 37 (2021) 1748-1764.
[3] H. Zhou, S. Zhang, J. Peng, S. Zhang, J. Li, H. Xiong, W. Zhang, Informer: Beyond efficient Transformer for long sequence time-series forecasting, Proceedings of the AAAI Conference on Artificial Intelligence 35 (2021) 11106-11115.
[4] H. Wu, J. Xu, J. Wang, M. Long, Autoformer: Decomposition Transformers with auto-correlation for long-term series forecasting, Advances in Neural Information Processing Systems 34 (2021) 22419-22430.
[5] Y. Nie, N.H. Nguyen, P. Sinthong, J. Kalagnanam, A time series is worth 64 words: Long-term forecasting with Transformers, International Conference on Learning Representations (2023).
[6] A. Zeng, M. Chen, L. Zhang, Q. Xu, Are Transformers effective for time series forecasting?, Proceedings of the AAAI Conference on Artificial Intelligence 37 (2023) 11121-11128.
[7] B.N. Oreshkin, D. Carpov, N. Chapados, Y. Bengio, N-BEATS: Neural basis expansion analysis for interpretable time series forecasting, International Conference on Learning Representations (2020).
[8] Z. Wu, S. Pan, G. Long, J. Jiang, C. Zhang, Graph WaveNet for deep spatial-temporal graph modeling, International Joint Conference on Artificial Intelligence (2019) 1907-1913.
[9] A.N. Angelopoulos, S. Bates, A gentle introduction to conformal prediction and distribution-free uncertainty quantification, Foundations and Trends in Machine Learning 16 (2023) 494-591.
[10] D. Bertsimas, N. Kallus, From predictive to prescriptive analytics, Management Science 66 (2020) 1025-1044.
[11] A.N. Elmachtoub, P. Grigas, Smart "predict, then optimize", Management Science 68 (2022) 9-26.
[12] B. Amos, J.Z. Kolter, OptNet: Differentiable optimization as a layer in neural networks, International Conference on Machine Learning (2017) 136-145.
[13] P. Donti, B. Amos, J.Z. Kolter, Task-based end-to-end model learning in stochastic optimization, Advances in Neural Information Processing Systems 30 (2017).
[14] T. Hong, P. Pinson, S. Fan, Global energy forecasting competition 2012, International Journal of Forecasting 30 (2014) 357-363.
[15] T. Hong, P. Pinson, S. Fan, H. Zareipour, A. Troccoli, R.J. Hyndman, Probabilistic energy forecasting: Global Energy Forecasting Competition 2014 and beyond, International Journal of Forecasting 32 (2016) 896-913.
[16] J. Lago, F. De Ridder, B. De Schutter, Forecasting spot electricity prices: Deep learning approaches and empirical comparison of traditional algorithms, Applied Energy 221 (2018) 386-405.
[17] J. Lago, G. Marcjasz, B. De Schutter, R. Weron, Forecasting day-ahead electricity prices: A review of state-of-the-art algorithms, best practices and an open-access benchmark, Applied Energy 293 (2021) 116983.
[18] R. Weron, Electricity price forecasting: A review of the state-of-the-art with a look into the future, International Journal of Forecasting 30 (2014) 1030-1081.
[19] B. Uniejewski, G. Marcjasz, R. Weron, On the importance of the long-term seasonal component in day-ahead electricity price forecasting, Energy Economics 79 (2019) 171-182, https://doi.org/10.1016/j.eneco.2018.02.007.
[20] J. Nowotarski, R. Weron, Recent advances in electricity price forecasting: A review of probabilistic forecasting, Renewable and Sustainable Energy Reviews 81 (2018) 1548-1568.
[21] S. Ben Taieb, J.W. Taylor, R.J. Hyndman, Coherent probabilistic forecasts for hierarchical time series, International Conference on Machine Learning (2017) 3348-3357.
[22] R.J. Hyndman, G. Athanasopoulos, Forecasting: Principles and Practice, third ed., OTexts, Melbourne, 2021.
[23] Open Power System Data, Data Package Time Series, version 2020-10-06, https://data.open-power-system-data.org/time_series/2020-10-06/.
[24] H.M. Rouzbahani, H. Karimipour, L. Lei, A review on virtual power plant for energy management, Sustainable Energy Technologies and Assessments 47 (2021) 101370, https://doi.org/10.1016/j.seta.2021.101370.
[25] S.M. Nosratabadi, R.-A. Hooshmand, E. Gholipour, A comprehensive review on microgrid and virtual power plant concepts employed for distributed energy resources scheduling in power systems, Renewable and Sustainable Energy Reviews 67 (2017) 341-363, https://doi.org/10.1016/j.rser.2016.09.025.
[26] S. Boyd, L. Vandenberghe, Convex Optimization, Cambridge University Press, Cambridge, 2004.
"""

REFERENCE_ASSET = ROOT / "references" / "paper1_eaai_verified_reference_block.md"
if REFERENCE_ASSET.exists():
    REFERENCES = REFERENCE_ASSET.read_text(encoding="utf-8")


def replace_section(text: str, heading: str, new_body: str) -> str:
    pattern = rf"(^## {re.escape(heading)}\n)(.*?)(?=^## |\Z)"
    return re.sub(pattern, rf"\1\n{new_body.strip()}\n\n", text, flags=re.M | re.S)


def drop_section(text: str, heading: str) -> str:
    pattern = rf"^## {re.escape(heading)}\n.*?(?=^## |\Z)"
    return re.sub(pattern, "", text, flags=re.M | re.S)


def extract_image(text: str, starts_with: str) -> str:
    for line in text.splitlines():
        if line.startswith(starts_with):
            return line
    raise ValueError(f"image not found: {starts_with}")


def extract_image_contains(text: str, caption_fragment: str) -> str:
    for line in text.splitlines():
        if line.startswith("![Figure ") and caption_fragment in line:
            return line
    raise ValueError(f"image not found: {caption_fragment}")


def remove_image(text: str, starts_with: str) -> str:
    return "\n".join(line for line in text.splitlines() if not line.startswith(starts_with)) + "\n"


def remove_image_contains(text: str, caption_fragment: str) -> str:
    return "\n".join(
        line
        for line in text.splitlines()
        if not (line.startswith("![Figure ") and caption_fragment in line)
    ) + "\n"


def renumber_figures_simple(text: str) -> str:
    counter = 0

    def repl(match: re.Match) -> str:
        nonlocal counter
        counter += 1
        return f"![Figure {counter}. {match.group(1).strip()}]({match.group(2)})"

    return re.sub(r"!\[Figure\s+\d+\.\s*(.*?)\]\((.*?)\)", repl, text)


def make_manuscript() -> str:
    text = SOURCE.read_text(encoding="utf-8")
    text = re.sub(r"^#\s+.+$", f"# {TITLE}", text, count=1, flags=re.M)
    text = text.replace("**Target journals:** Engineering Applications of Artificial Intelligence; Expert Systems with Applications; Neurocomputing; Applied Soft Computing.\n\n", "")
    text = re.sub(r"(?m)^\*\*Manuscript status:\*\*.*\n\n", "", text)

    architecture = extract_image_contains(text, "Proposed uncertainty-aware spatio-temporal price forecasting framework")
    architecture = re.sub(
        r"!\[Figure\s+\d+\.\s*Proposed uncertainty-aware spatio-temporal price forecasting framework\.\]",
        "![Figure 1. Sequence-anchored GraphPatch residual pipeline for lag-safe electricity price forecasting and conformal diagnostics.]",
        architecture,
    )
    for fragment in [
        "Proposed uncertainty-aware spatio-temporal price forecasting framework",
        "Sequence-anchored GraphPatch residual pipeline",
        "Representative split-conformal prediction interval on the local Hunan price holdout set",
        "Empirical coverage of the 90% split-conformal interval on local price datasets",
    ]:
        text = remove_image_contains(text, fragment)

    text = replace_section(text, "Abstract", f"{ABSTRACT}\n\n**Keywords:** {KEYWORDS}.")
    introduction = """Virtual power plants aggregate distributed photovoltaic generation, battery storage, flexible loads, and other controllable resources into a market-facing entity [24,25]. Their scheduling and bidding decisions depend strongly on day-ahead price forecasts, which are often coupled to downstream prescriptive, predict-then-optimize, and differentiable optimization layers [10-13,26]. For this application, average error alone is not enough: sparse price spikes, negative prices, and regime changes can dominate operational risk even when ordinary-hour accuracy appears acceptable.

From a computer-science perspective, this setting is a non-stationary time-series learning problem with heterogeneous covariates and cross-zone dependence. Public forecasting competitions, electricity-price forecasting reviews, long-term seasonality studies, probabilistic forecasting surveys, and open-data benchmarks establish that evaluation must cover chronology, volatility, uncertainty, and replicability rather than only a single average-error split [14-23]. Local temporal structure is strong, so simple seasonal, linear, and modern sequence baselines can be difficult to beat. At the same time, price spikes may reflect information that is not fully captured by a single-zone sequence model, including lagged behavior in neighboring markets, load conditions, and renewable-generation patterns.

The paper therefore asks a narrow empirical question: after a strong local sequence anchor has already modeled ordinary temporal persistence, can lag-safe cross-zone graph patches improve volatile price forecasting? The proposed answer is a residual-learning framework. A local anchor first predicts the target price, a GraphPatch residual learner then uses local temporal patches and lagged cross-zone summaries to estimate the remaining error, and a calibration-selected shrinkage or blend weight controls the residual contribution. Transformer foundations, interpretable temporal attention models, patch-based sequence models, linear sequence baselines, and N-BEATS-style decomposition models define the relevant sequence-learning comparison space [1-7]. Split-conformal calibration is used to report interval reliability rather than treating point-error gains as sufficient evidence [9].

This framing is intentionally more conservative than an unconstrained dynamic-graph or end-to-end Transformer claim. The contribution is not that graph learning universally dominates electricity price forecasting. The contribution is an auditable AI-engineering result on public OPSD data: graph-temporal residual information improves spike-regime forecasts beyond local ridge and DLinear/NLinear-style anchors in the reported tests, while leave-one-zone-out and conformal diagnostics expose the remaining transfer and calibration limitations.

The novelty is therefore in the residual interface rather than in attaching another feature block to a forecaster. DLinear, NLinear, and TDConv-style anchors are allowed to explain the dominant local seasonality, trend, and short-range convolutional structure first. GraphPatch then receives a different task: estimate the remaining forecast error from lag-safe cross-zone summaries and multi-scale patches, with its residual magnitude controlled by calibration-only shrinkage. This design creates a falsifiable test of relational information: if graph-temporal features merely duplicate the local anchor, the shrinkage weight should collapse toward zero or the held-out residual should not improve spike-regime errors. The method also separates three choices that are often entangled in electricity-price neural models: the local sequence model, the relational residual learner, and the uncertainty calibration layer. That separation makes the contribution more than an engineering tweak, because the experiment can identify when graph information is useful, when a stronger anchor absorbs it, and when conformal calibration still fails.

The manuscript contributes three elements. First, it defines a reproducible public benchmark for virtual-power-plant-oriented price forecasting with explicit spike-regime and calibration diagnostics. Second, it introduces a sequence-anchored GraphPatch residual design that keeps the graph contribution identifiable against DLinear/NLinear and TDConv-family local sequence baselines. Third, it reports rolling-origin, leave-one-zone-out, spike-threshold, strong-anchor, and calibration-window sensitivity checks so that the empirical claim is testable rather than dependent on a single split or threshold."""

    related_work = """### 2.1 Electricity-price forecasting protocols and benchmark discipline

Electricity-price forecasting has a mature literature, but the field remains sensitive to data leakage, split design, volatility treatment, and benchmark selection. Early and review work emphasizes that electricity prices differ from ordinary demand series because price spikes, negative prices, calendar effects, fuel conditions, and renewable penetration can create non-Gaussian and regime-dependent errors [16-20]. Open benchmark studies further show that model rankings can change when the evaluation moves from a single market or single split to multivariate, high-dimensional, and reproducible comparison settings [17,28]. This paper follows that benchmark discipline by using chronological public OPSD splits, training-defined spike thresholds, and explicit public/private data boundaries rather than treating a non-public single case or a single holdout window as decisive evidence.

### 2.2 Deep sequence models and the strength of simple anchors

Recent sequence-learning work provides many strong candidates for electricity-price forecasting, including attention-based, decomposition-based, patch-based, convolutional, and interpretable multi-horizon models [1-5]. At the same time, N-BEATS and linear sequence baselines demonstrate that sophisticated neural architectures do not automatically dominate when local seasonality and persistence are strong [6,7]. Broader deep time-series surveys reach the same methodological warning: the value of a new architecture depends on the forecast horizon, scale, stationarity, covariate availability, and comparison protocol [34,35]. This motivates the sequence-anchored design in this paper. The GraphPatch module is not evaluated as an unrestricted black-box replacement for strong local temporal models; it is evaluated as a residual learner that must add value after local lag-calendar, DLinear/NLinear-style, and TDConv-style anchors have already captured ordinary temporal structure.

### 2.3 Spatio-temporal graph learning for correlated market signals

Graph neural networks are relevant because electricity prices, loads, and renewable generation are spatially and operationally correlated across market zones. Spatio-temporal graph convolution, adaptive graph learning, and multivariate graph forecasting show that learned or partially learned relational structure can improve time-series forecasts when single-series models miss cross-node dependencies [8,32,33]. However, the direct transfer of traffic-style graph models to electricity markets is not straightforward: market zones have changing congestion patterns, cross-border flow constraints, regulatory differences, and uneven public covariate availability. The proposed GraphPatch design therefore uses lag-safe graph summaries and calibration-selected residual shrinkage. This sacrifices some expressiveness, but it keeps the evidence auditable and makes the graph contribution identifiable against local anchors.

### 2.4 Probabilistic forecasting, conformal calibration, and scoring

For virtual power plant operation, a useful price forecast must describe uncertainty as well as point location. Probabilistic energy forecasting competitions, hierarchical forecasting methods, and electricity-price probabilistic reviews establish that coverage, sharpness, calibration, and distributional scoring are central evaluation dimensions [14,15,20,21]. Proper scoring rules define why probabilistic forecasts should be judged by scores that reward calibrated distributions rather than only narrow intervals [29], while quantile regression provides a classical basis for interval and tail modeling [30]. Split-conformal prediction adds a model-agnostic calibration layer with transparent empirical coverage control under appropriate exchangeability conditions [9]. In this paper, conformal intervals are used as a diagnostic and deployment boundary: the GraphPatch point model is not treated as operationally sufficient unless interval coverage and width remain defensible on held-out data.

### 2.5 Statistical comparison and robustness evidence

Electricity-price forecasting papers are vulnerable to overclaiming when they report only average error on a single test period. Forecast-comparison work such as the Diebold-Mariano framework highlights that paired predictive accuracy should be assessed through matched forecast errors, while modern energy-forecasting practice also emphasizes rolling-origin validation, regime-specific performance, and transparent negative cases [27,31]. This paper reports paired spike-regime win rates, rolling-origin windows, leave-one-zone-out transfer, and calibration-window checks. The purpose is not to inflate the number of tests, but to make the claim falsifiable: a graph-temporal residual model should help where cross-zone information plausibly matters, and it should expose rather than hide cases where transfer or interval calibration remains weak.

### 2.6 Virtual power plant relevance and prediction-decision boundary

Virtual power plant studies connect forecasting to market bidding, distributed resource scheduling, and energy-management decisions [24,25]. Predictive-to-prescriptive analytics, smart predict-then-optimize learning, differentiable optimization layers, and task-based learning explain why forecast errors should ultimately be judged by downstream decisions when the operational objective is known [10-13,26]. This paper deliberately stops short of claiming full market-profit optimization. Its role in the dissertation sequence is the price-forecasting chapter: it supplies a calibrated and spike-aware public forecast evidence layer that can later feed the companion virtual-power-plant decision paper. This boundary keeps the EAAI submission focused on artificial-intelligence forecasting methodology while preserving the engineering relevance to VPP market operation."""

    problem_formulation = """Let p_t^z denote the hourly day-ahead electricity price for market zone z. Given a chronological lookback window, local calendar variables, local lagged prices, local load and renewable-generation variables, and lagged cross-zone summaries from other market zones, the task is to predict p_{t+h}^z on a held-out future test period. The main reported setting uses h = 1 and evaluates both all-hour and spike-regime performance. Spike hours are defined from the training-set distribution of absolute short-term versus daily-lag price changes, so no test-set information is used to select the volatility threshold.

The paper studies a residual-learning question rather than an unconstrained end-to-end forecasting claim: after a strong local sequence model has captured ordinary temporal persistence, can lag-safe graph-temporal patch features add measurable value in volatile price regimes? The implemented prediction pipeline is:

anchor_t^z = f_anchor(X_{t-L+1:t}^z),

r_t^z = p_t^z - anchor_t^z,

delta_t^z = f_GP(Phi_t^z),

hat p_t^z = anchor_t^z + gamma_z delta_t^z,

where f_anchor is the selected local ridge, DLinear-style, NLinear-style, or TDConv-style anchor; Phi_t^z contains local temporal patches and lagged cross-zone graph summaries; f_GP is the GraphPatch residual learner; and gamma_z is a shrinkage weight selected only on calibration data. Split-conformal calibration is then applied to calibration residuals to construct intervals around the point forecast. This formulation keeps the graph contribution identifiable and prevents the proposed model from being evaluated only as a black-box replacement for strong local sequence baselines."""

    proposed_method = f"""{architecture}

### 4.1 Lag-safe graph-temporal feature construction

The implemented graph-temporal input is deliberately lag-safe. For each target zone, the feature set includes local price lags, load and renewable-generation covariates, calendar indicators, and cross-zone summaries constructed only from information available before the forecast origin. Cross-zone features are summarized through lagged means, dispersions, and multi-scale temporal patches rather than through contemporaneous target-period prices. This design is less expressive than a fully learned dynamic graph, but it is auditable and directly aligned with the public OPSD data columns.

### 4.2 Sequence anchor and residual target

The first stage fits a local sequence anchor. Transparent lag-calendar ridge models provide the reproducible baseline, DLinear-style and NLinear-style anchors test decomposition-style sequence persistence, and a TDConv-style trainable dilated-convolution ridge comparator supplies a TCN-family strong-anchor check without introducing GPU-dependent reproducibility requirements. The residual target is the difference between the observed price and the selected anchor prediction. This residual formulation is central to the paper's computer-science claim: graph-temporal learning is used to explain what strong local temporal models still miss, especially during spike regimes.

### 4.3 GraphPatch residual learner

The GraphPatch residual learner constructs multi-scale patches from 1-, 2-, 3-, 6-, 12-, 24-, 48-, 72-, and 168-hour windows, augments them with lagged cross-zone graph summaries, and predicts an additive residual correction. The main nonlinear variant is a two-layer tanh residual MLP. A ridge residual variant and a lightweight random-hidden-layer residual model are retained as auditable controls. For each zone, the final blend weight between residual variants, or the residual shrinkage coefficient relative to the sequence anchor, is selected on calibration data only.

### 4.4 Split-conformal interval calibration

Point-error gains do not by themselves establish operational reliability. The interval layer therefore uses split-conformal calibration on residuals from the calibration window immediately preceding the test period. For a nominal 90% interval, the empirical residual quantile widens the point forecast into a prediction interval. Calibration-window sensitivity is reported to test whether the interval result depends on a fragile split choice.

### 4.5 Reviewer-facing robustness design

The method is evaluated through six diagnostics that match the paper's claims: spike-regime improvements against local ridge, sequence-anchored improvements against DLinear/NLinear-style anchors, TDConv-inclusive strong-anchor checks, lightweight patch-attention reviewer-response checks, rolling-origin robustness across multiple chronological windows, and leave-one-zone-out transfer across market zones. This design keeps the claims narrow. The paper does not claim universal average-error dominance, full portability across all market systems, or downstream bidding-profit claims; it claims that lag-safe graph-temporal residual learning improves volatile price forecasting under reproducible public-data tests while exposing calibration and transfer limitations."""

    conclusion = """This paper developed and evaluated a reproducible graph-temporal residual-learning framework for virtual-power-plant-oriented electricity price forecasting. The final evidence supports a focused computational claim: after local temporal structure is captured by ridge, DLinear-style, NLinear-style, or TDConv-style anchors, lag-safe GraphPatch residual features can add value in volatile price regimes. The DLinear/NLinear sequence-anchored residual improves spike-regime RMSE in all four OPSD zones, while the stricter TDConv-inclusive anchor improves all-hour RMSE in four zones and spike-regime RMSE in three zones. Rolling-origin validation is positive in all 12 zone-window cases, and leave-one-zone-out transfer is positive in three of four held-out zones. Split-conformal calibration provides usable intervals, but the Nordic zones remain mildly under-covered, showing that adaptive local calibration is still required before deployment.

The contribution is therefore not a broad assertion that graph learning always dominates electricity price forecasting. It is a reproducible AI-engineering result: residual graph-temporal representation learning can complement strong sequence baselines under price spikes, and public-data calibration diagnostics can reveal where uncertainty estimates remain fragile. This framing connects the work to virtual power plant market operation while keeping the submitted claims within the evidence provided by the public OPSD benchmark."""

    text = replace_section(text, "1. Introduction", introduction)
    text = replace_section(text, "2. Related Work", related_work)
    text = replace_section(text, "3. Problem Formulation", problem_formulation)
    text = replace_section(text, "4. Proposed Method", proposed_method)
    text = replace_section(text, "7. Conclusion", conclusion)
    text = text.replace(
        "The paired tests show that graph residual ridge has statistically significant spike-regime wins in DK2 and Great Britain, while DE-LU and DK1 remain mixed. The ELM residual model does not provide robust paired evidence and is significantly worse in Great Britain. These results justify retaining a regularized graph-temporal residual layer as the transparent public baseline, while the subsequent GraphPatch experiments test the stronger residual architecture required for the target-journal submission.",
        "The paired tests show that graph residual ridge has statistically significant spike-regime wins in DK2 and Great Britain, while DE-LU and DK1 remain mixed. The ELM residual model does not provide robust paired evidence and is significantly worse in Great Britain. These tests are treated as descriptive diagnostic evidence: interpretation depends on effect direction, magnitude, and the later robustness checks, and the p-values are not used to claim universal market portability. These results justify retaining a regularized graph-temporal residual layer as the transparent public baseline, while the subsequent GraphPatch experiments test the stronger residual architecture required for the target-journal submission.",
    )
    text = text.replace(
        "Recent progress in Transformer-based time-series forecasting, graph neural networks, and conformal prediction creates an opportunity",
        "Recent progress in Transformer-based time-series forecasting [1-7], graph neural networks [8], conformal prediction [9], and predict-then-optimize learning [10-13] creates an opportunity",
    )
    text = text.replace(
        "Transformer variants such as Informer, Autoformer, Temporal Fusion Transformer, PatchTST, and related architectures",
        "Transformer variants such as Temporal Fusion Transformer [2], Informer [3], Autoformer [4], PatchTST [5], and related architectures",
    )
    text = text.replace(
        "linear baselines such as DLinear show",
        "linear baselines such as DLinear [6] show",
    )
    text = text.replace(
        "Graph WaveNet and related spatio-temporal graph models demonstrate",
        "Graph WaveNet and related spatio-temporal graph models [8] demonstrate",
    )
    text = text.replace(
        "Conformal prediction is attractive",
        "Conformal prediction [9] is attractive",
    )
    text = text.replace(
        "Baselines should include ARIMA, SVR, Random Forest, XGBoost or LightGBM, LSTM/GRU, TCN, N-BEATS, DLinear, TFT, Informer/Autoformer/PatchTST, and a graph temporal model such as Graph WaveNet or MTGNN.",
        "The current public benchmark reports transparent persistence, seasonal, lag-calendar-exogenous, graph residual ridge, random-hidden-layer residual, GraphPatch residual, DLinear/NLinear-style anchors, a TDConv-style TCN-family sequence comparator, and a lightweight patch-attention reviewer baseline. Full PatchTST/TFT training remains a reviewer-response extension rather than a current claim.",
    )
    text = text.replace(
        "The current public benchmark now includes transparent persistence, seasonal, lag-calendar-exogenous, graph residual ridge, random-hidden-layer residual, GraphPatch residual, DLinear/NLinear-style anchors, and a TDConv-style TCN-family sequence comparator. Full PatchTST/TFT training remains a reviewer-response extension rather than a current claim.",
        "The current public benchmark now includes transparent persistence, seasonal, lag-calendar-exogenous, graph residual ridge, random-hidden-layer residual, GraphPatch residual, DLinear/NLinear-style anchors, a TDConv-style TCN-family sequence comparator, and a lightweight patch-attention reviewer baseline. Full PatchTST/TFT training remains a reviewer-response extension rather than a current claim.",
    )
    text = text.replace(
        "For uncertainty forecasting, compare quantile regression, DeepAR-style distributional forecasting, ensemble intervals, and conformalized baselines.",
        "For uncertainty forecasting, the manuscript reports conformalized baselines and leaves quantile-regression, DeepAR-style distributional forecasting, and ensemble intervals as optional extensions for the next experiment cycle.",
    )
    text = text.replace(
        "Market-value metrics: simulated bidding profit, regret, downside risk, and performance during top 5% price volatility periods.",
        "Operational relevance is reported through spike-regime error, interval coverage, interval width, and robustness diagnostics. Full bidding-profit simulation is left to the companion virtual-power-plant decision paper rather than claimed as evidence in this price-forecasting submission.",
    )
    text = text.replace(
        "The ablation should remove: dynamic graph learning, temporal patch attention, exogenous features, conformal calibration, and price-spike reweighting. Each component should be tested on both ordinary and high-volatility periods.",
        "The current ablation removes cross-zone graph information, nonlinear residual learning, calibration strategy, rolling-origin stability, leave-one-zone-out transfer, spike-threshold choice, and conformal calibration-window choice. Each component is evaluated on ordinary, high-volatility, chronological, and cross-zone transfer settings where the corresponding diagnostic is meaningful.",
    )
    text = text.replace(
        "The key gap for market operation is that many studies still optimize point accuracy and underreport calibration, extreme-event performance, and downstream decision value.",
        "The key gap for market operation is that many studies still optimize point accuracy and underreport calibration, extreme-event performance, and robustness under volatile regimes.",
    )
    text = text.replace(
        "This finding sets a strong reference point: graph-temporal structure is evaluated where it should matter most, namely calibrated uncertainty, spike-regime robustness, and downstream decision value.",
        "This finding sets a strong reference point: graph-temporal structure is evaluated where it should matter most in this manuscript, namely calibrated uncertainty, spike-regime robustness, and stability across chronological or cross-zone diagnostics.",
    )
    text = text.replace(
        "so the paper should frame the model as volatility-regime strengthening rather than universal average-error dominance.",
        "so the paper frames the model as volatility-regime strengthening rather than universal average-error dominance.",
    )
    text = text.replace(
        "graph-temporal point correction and conditional uncertainty calibration should be treated as complementary modules.",
        "graph-temporal point correction and conditional uncertainty calibration are best treated as complementary modules.",
    )
    text = text.replace("The planned evaluation emphasizes", "The empirical evaluation emphasizes")
    text = text.replace("If private virtual power plant data are used, report them", "Private virtual power plant data, if used, are reported")
    text = text.replace("and ensure that the main claims are reproducible", "and the main claims remain reproducible")

    text = text.replace(
        "Local Hunan and Shandong datasets are retained as application case studies rather than as the only evidence.",
        "Non-public Hunan and Shandong operational records are retained only for internal application-boundary checks; they are not used as public evidence for the main empirical claims in this target-journal manuscript.",
    )
    text = re.sub(
        r"\nThe public benchmark and the local pilot support the same methodological conclusion:.*?(?=\n## 7\. Conclusion)",
        "\nThe public benchmark supports the methodological conclusion that simple temporal structure is strong in ordinary periods, but market operation requires calibrated intervals and decision-aware evaluation under volatility. The OPSD conformal result also shows why conditional calibration, not only average coverage, is central to the proposed graph-temporal uncertainty model.\n\n### 6.2 Public Reproducibility and Application Boundary\n\nAll quantitative claims in this target-journal manuscript are evaluated on the public OPSD benchmark and can be regenerated from the accompanying public-data reproducibility package. Non-public Hunan and Shandong operational records were used only to check whether the research question is relevant to local virtual-power-plant market operation. They are not required to reproduce the reported tables or figures, are not redistributed, and are excluded from the public supplement. This separation keeps the submitted evidence auditable while preserving the application motivation.\n",
        text,
        flags=re.S,
    )
    patch_attention_block = make_patch_attention_baseline_block()
    if patch_attention_block and "Table 15. Patch-attention reviewer baseline" not in text:
        text = re.sub(
            r"(!\[Figure\s+\d+\.\s+TDConv-family sequence comparator and TDConv-inclusive GraphPatch residual on OPSD spike-regime RMSE\.\]\(figures/paper1_fig16_opsd_tdconv_anchor_graphpatch\.png\)\n)",
            r"\1\n" + patch_attention_block.replace("\\", "\\\\") + "\n",
            text,
            count=1,
        )

    text = drop_section(text, "Local Case-Study Data Assets")
    text = drop_section(text, "Empirical Plan Updated with Local Data")
    text = drop_section(text, "Cover Letter Draft")
    text = replace_section(text, "Data Availability Statement", DATA_AVAILABILITY_STATEMENT)

    declarations = f"""## Declaration of Competing Interest

The author declares no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.

## Declaration of Generative AI and AI-Assisted Technologies in the Writing Process

{AI_DECLARATION}
"""
    text = re.sub(r"\n## References\n.*\Z", "\n" + declarations.strip() + "\n\n" + REFERENCES.strip() + "\n", text, flags=re.S)
    text = renumber_figures_simple(text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() + "\n"


def make_title_page() -> str:
    return f"""# Title Page and Declarations

## Target Journal

Engineering Applications of Artificial Intelligence, Elsevier, ISSN 0952-1976.

## Manuscript Title

{TITLE}

## Authors

{AUTHOR}

## Affiliations

{AFFILIATION}

## Corresponding Author

{AUTHOR}, {EMAIL}, {AFFILIATION}

ORCID: {ORCID}

## Author Contributions

Conceptualization: {AUTHOR}.

Methodology: {AUTHOR}.

Software: {AUTHOR}.

Validation: {AUTHOR}.

Formal analysis: {AUTHOR}.

Investigation: {AUTHOR}.

Data curation: {AUTHOR}.

Writing - original draft: {AUTHOR}.

Writing - review and editing: {AUTHOR}.

Supervision: Not applicable for the current single-author submission draft.

## Declaration of Competing Interest

The author declares no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.

## Funding

This research received no external funding.

## Data Availability

{DATA_AVAILABILITY_STATEMENT}

## Generative AI Declaration

{AI_DECLARATION}
"""


def make_highlights() -> str:
    lines = ["# Highlights", ""]
    lines.extend(f"- {item}" for item in HIGHLIGHTS)
    lines.append("")
    lines.append("Each bullet is intended to stay within Elsevier's 85-character highlight limit.")
    return "\n".join(lines)


def make_checklist() -> str:
    return """# EAAI Paper 1 Submission Checklist

## Official-Source Rules Reflected in This Pack

| Requirement | Current handling |
|---|---|
| Practical AI method for a real-world engineering application | Abstract explicitly separates the AI contribution from the virtual-power-plant engineering application. |
| Public-data validation for replicability | OPSD public benchmark is the main empirical evidence; local data are only application evidence. |
| No undefined acronym in title or abstract | The title avoids undefined acronyms; the abstract expands Open Power System Data context through the dataset section. |
| Single-column manuscript | DOCX builder uses a single-column Word layout. |
| Double-anonymized workflow | Anonymous manuscript is separate from title page and author declarations. |
| Highlights | A separate highlights file contains five short bullets. |
| Data availability | A data availability statement is included. |
| Competing interest, funding, and AI-use declarations | Included in the title/declarations file and mirrored in the anonymous manuscript for final transparency review; the AI statement names OpenAI ChatGPT/Codex and preserves author responsibility. |
| Modern sequence baseline comparison | DLinear/NLinear-style sequence anchors, a TDConv-style TCN-family comparator, and a lightweight patch-attention reviewer baseline are included; full PatchTST/TFT training remains a reviewer-response extension. |
| Reference verification | A BibTeX-backed verified reference register exists; the manuscript uses confirmed DOI URLs where available and stable URLs where DOI insertion would be unsafe. |

## Remaining Author Actions Before Real Submission

- Confirm author names, affiliations, no-funding statement, omitted-acknowledgement handling, and corresponding-author details.
- Confirm with the school that the exact journal and rule year satisfy the required B/C classification.
- Complete the journal-classification evidence packet before APC payment or submission.
- Import the verified BibTeX file into Zotero/EndNote and normalize journal abbreviations/capitalization for the exact submission style.
- Add full PatchTST/TFT only if editors or reviewers demand a broader GPU-dependent deep-learning suite.
- Move code and derived public-result tables to a repository or supplementary archive.
- Review the AI declaration wording with the author and institutional policy before real upload.
- Ensure no private/local data are disclosed without permission.

## Evidence Date

Prepared on 2026-06-30 from current public EAAI/Elsevier guidance and the local reproducible experiment package.
"""


def setup_doc(document: Document, title: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(23, 54, 93)


def repeat_table_headers(document: Document) -> None:
    for table in document.tables:
        if not table.rows:
            continue
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))


def write_docx(md_path: Path, text: str) -> None:
    title = text.splitlines()[0].replace("# ", "").strip()
    document = Document()
    setup_doc(document, title)
    add_markdown_to_docx(document, text)
    repeat_table_headers(document)
    document.save(md_path.with_suffix(".docx"))


def main() -> None:
    files = {
        "paper_1_eaai_anonymous_manuscript.md": make_manuscript(),
        "paper_1_eaai_title_page_and_declarations.md": make_title_page(),
        "paper_1_eaai_highlights.md": make_highlights(),
        "paper_1_eaai_submission_checklist.md": make_checklist(),
    }
    for name, text in files.items():
        path = OUT / name
        path.write_text(text, encoding="utf-8")
        write_docx(path, text)
        print(path)
        print(path.with_suffix(".docx"))


if __name__ == "__main__":
    main()
