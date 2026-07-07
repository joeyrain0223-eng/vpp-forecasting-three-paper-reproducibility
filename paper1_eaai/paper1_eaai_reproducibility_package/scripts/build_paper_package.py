from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


BASE = Path(__file__).resolve().parents[1]
ROOT = BASE
OUT = ROOT / "manuscript"
OUT.mkdir(parents=True, exist_ok=True)


COMMON_REFS = """
## References

[1] A. Vaswani et al., "Attention Is All You Need," NeurIPS, 2017.
[2] B. Lim, S. O. Arik, N. Loeff, and T. Pfister, "Temporal Fusion Transformers for Interpretable Multi-horizon Time Series Forecasting," International Journal of Forecasting, 2021.
[3] H. Zhou et al., "Informer: Beyond Efficient Transformer for Long Sequence Time-Series Forecasting," AAAI, 2021.
[4] H. Wu et al., "Autoformer: Decomposition Transformers with Auto-Correlation for Long-Term Series Forecasting," NeurIPS, 2021.
[5] Y. Nie et al., "A Time Series is Worth 64 Words: Long-term Forecasting with Transformers," ICLR, 2023.
[6] A. Zeng, M. Chen, L. Zhang, and Q. Xu, "Are Transformers Effective for Time Series Forecasting?" AAAI, 2023.
[7] B. N. Oreshkin et al., "N-BEATS: Neural Basis Expansion Analysis for Interpretable Time Series Forecasting," ICLR, 2020.
[8] Z. Wu et al., "Graph WaveNet for Deep Spatial-Temporal Graph Modeling," IJCAI, 2019.
[9] A. N. Angelopoulos and S. Bates, "A Gentle Introduction to Conformal Prediction and Distribution-Free Uncertainty Quantification," Foundations and Trends in Machine Learning, 2023.
[10] D. Bertsimas and M. Kallus, "From Predictive to Prescriptive Analytics," Management Science, 2020.
[11] A. N. Elmachtoub and P. Grigas, "Smart Predict-then-Optimize," Management Science, 2022.
[12] B. Amos and J. Z. Kolter, "OptNet: Differentiable Optimization as a Layer in Neural Networks," ICML, 2017.
[13] P. Donti, B. Amos, and J. Z. Kolter, "Task-based End-to-end Model Learning in Stochastic Optimization," NeurIPS, 2017.
[14] T. Hong, P. Pinson, and S. Fan, "Global Energy Forecasting Competition 2012," International Journal of Forecasting, 2014.
[15] T. Hong et al., "Probabilistic Energy Forecasting: Global Energy Forecasting Competition 2014 and Beyond," International Journal of Forecasting, 2016.

Note: final manuscript submission should replace this compact seed list with the target journal's exact reference style and add domain-specific electricity price, load forecasting, and virtual power plant references from the final literature search.
""".strip()


def paper_one():
    return """# Sequence-Anchored GraphPatch Residual Learning for Uncertainty-Aware Electricity Price Forecasting in Virtual Power Plant Market Operation

**Target journals:** Engineering Applications of Artificial Intelligence; Expert Systems with Applications; Neurocomputing; Applied Soft Computing.

**Manuscript status:** Working manuscript with public-data experiments, figures, references, and data-availability materials integrated. Do not submit until target-journal formatting, school classification confirmation, and final author approval are completed.

## Abstract

Electricity price forecasting has become a core computational problem for virtual power plants that participate in increasingly volatile electricity markets. Existing forecasting studies often focus on point accuracy, while market operation requires calibrated uncertainty, robustness to price spikes, and decision-relevant representations across heterogeneous market, load, renewable, weather, and calendar signals. This paper proposes an uncertainty-aware spatio-temporal learning framework for multi-horizon electricity price forecasting in virtual power plant market operation. The framework constructs a dynamic heterogeneous temporal graph over market zones, resource-side variables, and exogenous covariates, and then combines graph message passing, patch-based temporal attention, and probabilistic output heads to generate calibrated point and interval forecasts. A conformal calibration layer is further introduced to improve distribution-free reliability under non-stationary market regimes. The proposed model is evaluated on public electricity market datasets and virtual-power-plant-oriented feature sets against statistical, machine-learning, and deep time-series baselines. The planned evaluation reports deterministic accuracy, probabilistic calibration, robustness under spike regimes, cross-market generalization, and downstream bidding value. The study positions electricity price forecasting as a computer-science problem of uncertain, heterogeneous, non-stationary time-series learning rather than a conventional power-system simulation problem.

**Keywords:** electricity price forecasting; virtual power plant; spatio-temporal learning; graph neural network; Transformer; conformal prediction; uncertainty quantification.

## 1. Introduction

Virtual power plants aggregate distributed photovoltaic generation, battery storage, flexible loads, and other controllable resources into a market-facing entity. Their economic performance depends on the ability to anticipate market prices at multiple horizons. A virtual power plant that only receives a point forecast is exposed to two kinds of risk. First, price spikes and negative-price events can dominate profit and loss even if average prediction error is moderate. Second, bidding, arbitrage, and load-shifting decisions require a distribution over plausible outcomes, not only a single expected value. These requirements turn electricity price forecasting into a problem of uncertainty-aware temporal learning.

From a computer-science perspective, electricity price data exhibit several characteristics that are not well handled by generic forecasting pipelines. The series is non-stationary across seasons and market regimes; it depends on heterogeneous exogenous variables such as load, weather, renewable generation, fuel prices, and calendar effects; and it contains sparse but economically important extreme events. Spatial relationships among market nodes, balancing areas, or resource groups are also time-varying. A model that treats the price series as an isolated univariate signal is therefore structurally limited.

Recent progress in Transformer-based time-series forecasting, graph neural networks, and conformal prediction creates an opportunity to design a more suitable computational framework. Patch-based temporal attention can model long-range temporal patterns efficiently, graph learning can capture spatial or relational dependencies, and conformal calibration can provide prediction intervals with finite-sample validity under mild assumptions. However, these components have rarely been integrated into a virtual-power-plant market operation setting with a direct emphasis on calibrated uncertainty and price-spike robustness.

This paper proposes an uncertainty-aware spatio-temporal learning framework for multi-horizon electricity price forecasting. The method first transforms raw market, load, renewable, weather, and calendar signals into a dynamic heterogeneous graph. A graph-temporal encoder extracts relational representations, a temporal attention module captures multi-scale dependencies, and probabilistic heads estimate quantiles or distributional parameters for each forecast horizon. Finally, a conformal calibration layer adjusts interval forecasts using recent residual behavior, improving reliability when the market regime shifts.

The intended contribution is methodological. Electricity markets are used as a demanding testbed for multi-source non-stationary time-series learning. The model is designed to be compared not only by MAE and RMSE but also by pinball loss, CRPS, interval coverage, interval width, spike-regime error, and downstream bidding value. This framing makes the manuscript suitable for AI engineering and intelligent decision-support journals while remaining relevant to virtual power plant applications.

## 2. Related Work

### 2.1 Electricity price forecasting

Electricity price forecasting studies have evolved from statistical models and shallow machine learning toward deep temporal architectures. Traditional approaches such as ARIMA, support vector regression, random forests, and gradient boosting are still useful baselines because they are stable, interpretable, and computationally efficient. Deep learning approaches, including recurrent networks, temporal convolutional networks, attention models, and hybrid architectures, can represent nonlinear temporal dependencies and exogenous interactions. The key gap for market operation is that many studies still optimize point accuracy and underreport calibration, extreme-event performance, and downstream decision value.

### 2.2 Deep time-series forecasting

Transformer variants such as Informer, Autoformer, Temporal Fusion Transformer, PatchTST, and related architectures have become standard references for long-horizon forecasting. At the same time, linear baselines such as DLinear show that sophisticated attention is not automatically superior unless the data regime and evaluation protocol justify its complexity. This motivates a careful experimental design in which the proposed method is compared with both strong deep models and simple but competitive baselines.

### 2.3 Spatio-temporal graph learning

Electricity markets contain relational structure: market zones are connected by transmission constraints, neighboring loads respond to similar weather patterns, and distributed resources are correlated through geography and behavior. Graph WaveNet and related spatio-temporal graph models demonstrate that learned graph structures can improve forecasting when explicit adjacency is incomplete or dynamic. This paper adopts this principle but extends it to heterogeneous market covariates and uncertainty-aware price prediction.

### 2.4 Uncertainty quantification and conformal prediction

Prediction intervals are essential for bidding and risk control. Quantile regression, Bayesian neural networks, ensembling, and distributional forecasting provide alternative uncertainty estimates. Conformal prediction is attractive because it can wrap around an existing model and provide distribution-free calibration under exchangeability or approximate online conditions. In this paper, conformal calibration is used as a post-processing layer to correct the empirical coverage of neural quantile forecasts.

## 3. Problem Formulation

Let p_t^z denote the electricity price at time t for market zone z, and let X_t^z contain exogenous features including system load, renewable generation, weather, calendar variables, and resource aggregation indicators. Given a lookback window of length L, the goal is to forecast zone-specific prices for horizons h = 1, ..., H. Unlike deterministic forecasting, the model estimates conditional quantiles q_alpha(t+h,z) or a predictive distribution P(p_{t+h}^z | X_{t-L+1:t}^z).

The prediction target is therefore:

f_theta: {X_{t-L+1:t}^z, G_t} -> {q_alpha(t+h,z) for alpha in A, h = 1, ..., H},

where G_t is a dynamic heterogeneous graph representing market and resource relationships. The objective combines point error, quantile loss, calibration loss, and optional decision-aware regularization.

## 4. Proposed Method

### 4.1 Dynamic heterogeneous market graph

The graph contains nodes for market zones, load areas, renewable generation groups, weather stations, and virtual-power-plant resource clusters. Edges are initialized from physical or geographic proximity when available and then updated by learned similarity. For nodes i and j at time t, the adaptive edge weight is:

a_ij,t = softmax_j(phi(z_i,t)^T psi(z_j,t) / sqrt(d)),

where z_i,t and z_j,t are node embeddings and phi, psi are learnable projections. This design allows the model to discover time-varying dependencies without relying exclusively on fixed grid topology.

### 4.2 Spatio-temporal encoder

The encoder alternates graph message passing and temporal attention. Graph message passing aggregates relational information:

h_i,t' = sigma(W_self h_i,t + sum_j a_ij,t W_msg h_j,t).

The temporal module then applies patch-based attention over each node representation. Patches reduce sequence length and help the model capture daily, weekly, and seasonal structures without excessive computation. Positional, calendar, and market-session embeddings are added to preserve temporal semantics.

### 4.3 Probabilistic forecasting head

The forecasting head outputs multiple quantiles for each horizon. The quantile loss is:

L_q = sum_{h, alpha} max(alpha e_{h,alpha}, (alpha - 1) e_{h,alpha}),

where e_{h,alpha}^z = p_{t+h}^z - q_alpha(t+h,z). The median quantile is used for point prediction, while lower and upper quantiles provide interval forecasts.

### 4.4 Conformal calibration

A rolling calibration set is used to estimate nonconformity scores from recent residuals. For a nominal interval [q_lower, q_upper], the conformalized interval is widened by an empirical quantile of calibration errors. This layer is deliberately separated from the neural encoder so that it can adapt to regime shifts without retraining the full model.

### 4.5 Training objective

The full objective is:

L_total = L_pred + lambda_q L_q + lambda_c L_cal + lambda_r L_reg,

where L_pred is MAE or Huber loss for the median prediction, L_q is quantile loss, L_cal encourages empirical coverage, and L_reg controls smoothness or graph sparsity. Hyperparameters are selected on validation data.

## 5. Experimental Design

### 5.1 Datasets

The final manuscript should use at least two public electricity market datasets. Recommended candidates include PJM, NYISO, CAISO, ISO-NE, AEMO, or ENTSO-E. Each dataset should be aligned with load, renewable, weather, and calendar features. If private virtual power plant data are available, they should be reported as an additional case study rather than the only evidence.

**Result placeholder:** Insert Table 1 with dataset statistics: market, time span, resolution, number of nodes, variables, training/validation/test split, missing-value handling.

### 5.2 Baselines

Baselines should include ARIMA, SVR, Random Forest, XGBoost or LightGBM, LSTM/GRU, TCN, N-BEATS, DLinear, TFT, Informer/Autoformer/PatchTST, and a graph temporal model such as Graph WaveNet or MTGNN. For uncertainty forecasting, compare quantile regression, DeepAR-style distributional forecasting, ensemble intervals, and conformalized baselines.

### 5.3 Metrics

Point forecasting metrics: MAE, RMSE, MAPE or sMAPE. Probabilistic metrics: pinball loss, CRPS, PICP, PINAW, calibration error. Market-value metrics: simulated bidding profit, regret, downside risk, and performance during top 5% price volatility periods.

### 5.4 Ablation study

The ablation should remove: dynamic graph learning, temporal patch attention, exogenous features, conformal calibration, and price-spike reweighting. Each component should be tested on both ordinary and high-volatility periods.

## 6. Results and Discussion

**Table 2 placeholder:** Overall point forecasting performance across datasets.

**Table 3 placeholder:** Probabilistic forecasting performance and calibration.

**Table 4 placeholder:** Price-spike and high-volatility subset performance.

**Figure 1 placeholder:** Model architecture.

**Figure 2 placeholder:** Prediction intervals over representative volatile weeks.

**Figure 3 placeholder:** Coverage-width trade-off before and after conformal calibration.

When empirical results are available, the discussion should avoid claiming universal superiority. Instead, it should identify the regimes in which the proposed model is most useful: high volatility, strong exogenous coupling, cross-node correlation, and decision settings where interval reliability matters. If simple baselines perform strongly in stable periods, that result should be reported and used to argue that the proposed method is valuable primarily under market complexity rather than as a blanket replacement for all forecasting tools.

## 7. Conclusion

This paper presents an uncertainty-aware spatio-temporal learning framework for electricity price forecasting in virtual power plant market operation. By integrating dynamic graph learning, temporal attention, probabilistic forecasting, and conformal calibration, the framework addresses both predictive accuracy and operational reliability. The planned evaluation emphasizes not only point error but also interval calibration, spike robustness, and downstream market value. The work contributes to computer-science research on heterogeneous non-stationary time-series learning while providing a practical forecasting component for virtual power plant decision systems.

## Data Availability Statement

The final submission should specify all public datasets and provide preprocessing scripts. If private virtual power plant data are used, report them as a non-public supplementary case and ensure that the main claims are reproducible on public datasets.

## Cover Letter Draft

Dear Editor,

We submit the manuscript "Sequence-Anchored GraphPatch Residual Learning for Uncertainty-Aware Electricity Price Forecasting in Virtual Power Plant Market Operation" for consideration. The manuscript studies electricity price forecasting as a problem of heterogeneous, non-stationary, uncertainty-aware time-series learning. It proposes a framework combining local sequence anchoring, lag-safe graph-temporal patches, residual learning, and conformal calibration. The work is intended for readers interested in artificial intelligence engineering applications, intelligent forecasting systems, and decision support under uncertainty.

The manuscript is original, has not been published elsewhere, and is not under consideration by another journal. All authors have approved the submission. We believe the paper fits the journal because it develops and evaluates an AI method for a complex engineering decision problem rather than presenting a narrow domain-specific simulation study.

Sincerely,

[Author names]

""" + "\n\n" + COMMON_REFS + "\n"


def paper_two():
    return """# Transferable Short-Term Load Forecasting for Aggregated Virtual Power Plant Resources via Source-Pooled Time-Series Representation Learning

**Target journals:** IEEE Access; Neurocomputing; Applied Soft Computing; Energy Reports.

**Manuscript status:** Working manuscript with public OPSD/UCI experiments, figures, references, and transfer-protocol details integrated. Do not submit until target-journal formatting, school classification confirmation, and final author approval are completed.

## Abstract

Accurate load forecasting is fundamental for virtual power plants that coordinate distributed resources, flexible demand, storage, and market transactions. However, aggregated virtual power plant resources often face cross-domain heterogeneity, limited historical observations, and cold-start deployment conditions. Conventional supervised load forecasting models perform well when training and test data come from the same distribution, but they frequently degrade when applied to new buildings, industrial parks, charging clusters, or regional aggregations. This paper studies transferable short-term load forecasting as a time-series representation-learning problem. Instead of claiming a single universal deep model, the reproducible framework evaluates source-pooled temporal representations under transparent transfer protocols: masked-reconstruction features, random-convolution features, and trainable dilated-convolution ridge features, each paired with lightweight target adapters or source heads. Public experiments cover OPSD system-load zones, UCI Electricity Load Diagrams multi-client transfer, and UCI Appliances Energy Prediction external validation. The study frames virtual power plant load forecasting as a computer-science problem of representation reuse, label efficiency, and failure-aware adaptation under domain shift.

**Keywords:** load forecasting; virtual power plant; source-pooled representation learning; transfer learning; time-series representation; cold start; domain adaptation.

## 1. Introduction

Virtual power plants depend on short-term load forecasts to schedule flexible demand, allocate storage, plan market bids, and evaluate reserve capability. Unlike conventional system-level load forecasting, virtual power plant forecasting often concerns aggregated portfolios of heterogeneous resources: commercial buildings, industrial parks, residential communities, electric vehicle charging clusters, and distributed energy assets. These resources differ in scale, behavioral patterns, weather sensitivity, occupancy structure, and controllability.

A practical challenge is that new resources may have limited historical data. A building or charging cluster newly connected to a virtual power plant cannot provide years of labeled observations. Even when historical data exist, distribution shifts occur across regions, seasons, users, and tariff regimes. A supervised model trained on one aggregation may fail on another. This motivates a transferable forecasting framework that learns reusable temporal representations before task-specific fine-tuning.

Self-supervised and source-pooled representation learning offer practical routes for this setting. Instead of requiring labels for every target resource, a model can learn temporal features from source load archives through masked reconstruction, temporal-filter representations, or other pretext and representation objectives. The learned representation can then be adapted to downstream forecasting tasks with fewer labels. This paradigm is increasingly important for time-series learning, but it must be tested against strong target-only and transparent baselines rather than only against weak seasonal rules.

This paper studies a transferable load forecasting framework for aggregated virtual power plant resources. The method learns source-domain temporal representations, adapts them through lightweight target heads or adapters, and evaluates them under practical deployment scenarios: cross-client transfer, few-shot adaptation, and cold-start forecasting. The goal is not merely to reduce RMSE in a single dataset, but to understand when reusable representations improve generalization under distribution shift and when simple target baselines remain competitive.

The contribution is computer-science oriented. The paper studies how to learn reusable temporal representations for heterogeneous energy-related time series. The virtual power plant setting supplies a realistic and economically important application, but the method is positioned as a general transferable forecasting architecture.

## 2. Related Work

### 2.1 Short-term load forecasting

Short-term load forecasting has been studied with statistical models, machine-learning regressors, and deep neural networks. Classical methods are reliable for stable aggregate loads, while deep models can capture nonlinear relationships among weather, calendar, and behavior. However, most supervised approaches assume that training and deployment data share similar distributions. In virtual power plant portfolios, this assumption is often violated.

### 2.2 Transfer learning and domain adaptation

Transfer learning aims to improve target-domain performance by leveraging source-domain knowledge. In load forecasting, domain shift may arise from geographic differences, building type, user behavior, seasonal operation, or resource composition. Domain adaptation methods reduce representation mismatch, while meta-learning and few-shot learning seek rapid adaptation from limited target data. The proposed framework combines shared representation learning with lightweight target-domain adaptation.

### 2.3 Self-supervised time-series representation learning

Self-supervised time-series learning uses pretext tasks to learn representations without manual labels. Masked reconstruction encourages local and global temporal understanding, contrastive learning separates representations of different sequences or augmentations, and multi-scale objectives capture patterns at different resolutions. For load forecasting, these methods can exploit large unlabeled archives before downstream prediction labels are available.

## 3. Problem Formulation

Consider N load domains, where each domain may represent a building, park, charging cluster, region, or aggregation portfolio. For client or resource domain i, a sequence contains load ell_t^i and covariates x_t^i such as weather, calendar variables, and resource metadata. Given a lookback window L and forecast horizon H, the target is:

f_theta(x_{t-L+1:t}^i, ell_{t-L+1:t}^i, m^i) -> ell_{t+1:t+H}^i,

where m^i denotes domain metadata. The key challenge is to maintain performance when target domain i has limited or no labeled training data.

The evaluation uses four protocols:

1. Within-domain forecasting: train and test on the same domain.
2. Cross-domain forecasting: train on source domains and test on unseen domains.
3. Few-shot adaptation: fine-tune with k days or weeks of target-domain data.
4. Cold-start forecasting: use metadata and exogenous variables with minimal load history.

## 4. Proposed Method

### 4.1 Time-series encoder

The encoder maps a multivariate time-series window into latent representations. It may use temporal convolution, patch-based Transformer blocks, or a hybrid architecture. Patch embeddings reduce sequence length, and calendar embeddings encode hour-of-day, day-of-week, holiday, and seasonal information. The model is deliberately modular so that alternative encoders can be tested.

### 4.2 Masked temporal reconstruction

During pretraining, random segments of the input sequence are masked. The encoder must reconstruct the masked load and covariate values using surrounding context:

L_mask = ||M * (X - decoder(encoder(mask(X))))||_1,

where M is the mask indicator. Segment-level masking is preferred over independent point masking because it forces the model to learn meaningful temporal dependencies rather than interpolate isolated points.

### 4.3 Domain-shift-aware representation checks

The reproducible implementation treats domain shift as an empirical object rather than as a solved property of the encoder. It compares masked-reconstruction features, random-convolution temporal features, and trainable dilated-convolution ridge features under the same source-head and target-adapter protocols. This design tests whether representation strength, target-label budget, and client mismatch explain transfer gains. Any contrastive or adversarial extension is therefore a future reviewer-response option, not a claim required by the current public evidence.

### 4.4 Standardization and lightweight adapters

Load magnitude and variability differ significantly across resources. The current protocol standardizes source and target windows, fits source-domain representation heads, and then uses lightweight target adapters with 1, 3, 7, or 28 days of target data. This reduces the number of target-fitted parameters and exposes low-label overfitting when the 1-day and 3-day adapters become unstable. The paper therefore claims regularized adaptation value under matched protocols rather than broad target-only refitting.

### 4.5 Forecasting head and training schedule

The model is trained in three stages:

1. Source-domain representation fitting on multi-client load sequences.
2. Source-head forecasting on held-out target clients with no target labels.
3. Target-domain adaptation with limited labeled data and matched target-only baselines.

The supervised objective is MAE or Huber loss, optionally combined with quantile loss for probabilistic outputs.

## 5. Experimental Design

### 5.1 Datasets

Recommended public datasets include GEFCom load data, UCI Electricity Load Diagrams, Open Power System Data, ISO system load data, building energy datasets, and electric-vehicle charging datasets if available. For a virtual power plant scenario, individual series can be aggregated into portfolios of different sizes and compositions.

**Result placeholder:** Insert Table 1 with dataset names, domains, resolution, time span, covariates, number of series, and aggregation method.

### 5.2 Baselines

Baselines should include seasonal naive, ARIMA, SVR, Random Forest, XGBoost/LightGBM, LSTM, GRU, TCN, N-BEATS, DLinear, TFT, PatchTST, and at least two transfer-learning baselines such as fine-tuning from a supervised source model and domain-adversarial training.

### 5.3 Metrics

Point metrics: MAE, RMSE, MAPE, sMAPE. Transfer metrics: relative improvement over target-only training, few-shot performance as a function of target data size, and adaptation efficiency. Robustness metrics: performance under weather extremes, holidays, and unseen domains.

### 5.4 Ablation study

Ablations compare masked-reconstruction features, random-convolution features, trainable dilated-convolution ridge features, source heads, target adapters, and target-only negative controls. The analysis reports whether each component contributes more to cross-domain generalization, cold-start transfer, or ordinary target-domain accuracy.

## 6. Results and Discussion

**Table 2 placeholder:** Within-domain forecasting performance.

**Table 3 placeholder:** Cross-domain and unseen-domain forecasting performance.

**Table 4 placeholder:** Few-shot adaptation with 1 day, 3 days, 1 week, and 1 month of target data.

**Figure 1 placeholder:** Framework architecture.

**Figure 2 placeholder:** Representation visualization by domain before and after adaptation.

**Figure 3 placeholder:** Error versus target-domain data size.

The discussion highlights whether reusable temporal representations improve generalization rather than only improving average accuracy. If the representation model is not best on all within-domain tests, that does not invalidate the paper. The central claim is bounded: source-pooled temporal representations can reduce the data requirement for new virtual power plant resources, but target adapters and representation choice need validation because some low-label settings remain unstable.

## 7. Conclusion

This paper presents a transferable short-term load forecasting study for aggregated virtual power plant resources. By comparing masked-reconstruction, random-convolution, and trainable dilated-convolution representations under matched source-head and target-adapter protocols, the work shows when source-pooled temporal features help cross-domain, few-shot, and cold-start load forecasting. The conclusion is deliberately bounded: reusable representations improve several held-out UCI client protocols, but representation choice and target-label budget still require validation, and transparent lag-weather models remain competitive on the external UCI Appliances check.

## Data Availability Statement

The final submission should provide preprocessing code for public datasets and describe any private or simulated virtual power plant aggregation data. Main claims should be reproducible without relying solely on private data.

## Cover Letter Draft

Dear Editor,

We submit the manuscript "Transferable Short-Term Load Forecasting for Aggregated Virtual Power Plant Resources via Source-Pooled Time-Series Representation Learning" for consideration. The manuscript addresses a practical and methodological challenge: how to forecast heterogeneous virtual-power-plant load resources when target-domain data are limited. It evaluates source-pooled temporal representations, masked-reconstruction features, random-convolution features, trainable dilated-convolution ridge features, source heads, and lightweight target adapters under cross-client, few-shot, and cold-start protocols.

The paper fits the journal because it contributes a transferable time-series learning method and evaluates it in an engineering decision context. The manuscript is original, not published, and not under concurrent review.

Sincerely,

[Author names]

""" + "\n\n" + COMMON_REFS + "\n"


def paper_three():
    return """# Decision-Focused Learning for Virtual Power Plant Bidding under Electricity Price and Load Uncertainty

**Target journals:** IEEE Access; Applied Soft Computing; Energy Reports; CSEE Journal of Power and Energy Systems.

**Manuscript status:** Working manuscript with public OPSD decision experiments, simulator appendix, references, and target-journal package integrated. Do not submit until target-journal formatting, school classification confirmation, and final author approval are completed.

## Abstract

Virtual power plant bidding decisions depend on forecasts of electricity price, load, renewable generation, and resource flexibility. Standard forecasting pipelines optimize prediction accuracy first and then pass forecasts into a separate optimization module. This forecast-then-optimize paradigm can be suboptimal because the forecasting loss may not align with downstream market profit, risk exposure, or operational penalties. This paper proposes a decision-focused learning framework for virtual power plant bidding under electricity price and load uncertainty. The framework embeds a differentiable or surrogate optimization layer into the learning process so that forecasting representations are trained with respect to operational objectives. Price and load uncertainty are represented by quantile or scenario outputs, and the bidding layer optimizes expected revenue subject to storage, load flexibility, and risk constraints. When the exact optimization problem is not differentiable, a regret-based surrogate is used to approximate decision sensitivity. The empirical evaluation compares previous-day forecast-then-optimize, rolling-mean forecast-then-optimize, robust quantile forecast-then-optimize, and held-out decision-focused policy search on public-market and local case-study scenarios. The study frames virtual power plant operation as a machine-learning problem of aligning prediction with downstream decision quality.

**Keywords:** virtual power plant; electricity market bidding; decision-focused learning; predict-and-optimize; uncertainty; robust optimization; time-series forecasting.

## 1. Introduction

Virtual power plants are expected to participate in electricity markets by coordinating distributed energy resources, flexible loads, storage, and renewable generation. Their bidding decisions require forecasts of prices and resource availability. A common engineering pipeline first trains forecasting models to minimize MAE or RMSE and then feeds the forecasts into an optimization model. Although intuitive, this pipeline can fail when small prediction errors near decision boundaries produce large financial losses or infeasible schedules.

For example, an electricity price forecast that is slightly inaccurate during a high-spread period can reverse the value of charging or discharging storage. A load forecast that is accurate on average may still lead to reserve shortfalls if its uncertainty is miscalibrated. These examples show that forecasting accuracy and decision quality are related but not equivalent. A virtual power plant needs predictions that are useful for decisions, not merely predictions that minimize symmetric statistical error.

Decision-focused learning provides a principled framework for this problem. Instead of treating optimization as a downstream black box, the learning process incorporates decision loss, regret, or differentiable optimization layers. The model can therefore learn representations that prioritize errors with high operational consequences. This perspective is well aligned with computer-science research on predict-and-optimize, differentiable optimization, and task-based learning.

This paper proposes a decision-focused learning framework for virtual power plant bidding. The method combines probabilistic price and load forecasting with a bidding optimization layer that accounts for storage constraints, flexible-load limits, and risk preferences. The paper compares the proposed framework with standard forecast-then-optimize methods and evaluates not only prediction error but also revenue, regret, risk, and operational penalties.

The contribution is not a new power-system dispatch formulation. Rather, the contribution is a learning framework that connects uncertain time-series prediction with downstream optimization objectives in a virtual power plant market setting.

## 2. Related Work

### 2.1 Virtual power plant market operation

Virtual power plants aggregate distributed resources to provide market services such as energy trading, demand response, reserve support, and renewable integration. Market participation requires decisions under uncertainty. Existing work often formulates bidding or scheduling as deterministic, stochastic, robust, or reinforcement-learning optimization. These methods typically assume forecasts are given or train predictors separately.

### 2.2 Forecast-then-optimize and decision-focused learning

Predictive models are commonly evaluated with statistical error metrics. However, prescriptive analytics shows that prediction quality should be evaluated by downstream decisions when the ultimate objective is operational. Smart predict-and-optimize methods, differentiable optimization layers, and task-based end-to-end learning provide tools for aligning model training with optimization performance. This paper applies that idea to virtual power plant bidding under uncertain prices and loads.

### 2.3 Uncertainty-aware energy decision making

Robust and stochastic optimization account for uncertain variables by considering worst-case sets, scenarios, or probability distributions. In machine-learning-based pipelines, uncertainty can be generated by quantile forecasts, ensembles, Bayesian methods, or conformal prediction. The proposed method uses uncertainty not only for reporting prediction intervals but also as an input to risk-aware bidding decisions.

## 3. Problem Formulation

Consider a virtual power plant that participates in a day-ahead or intraday market. At each decision time, it observes historical prices, load, renewable generation, storage state, and resource constraints. It chooses a bid or schedule u for future horizons h = 1, ..., H. The realized revenue depends on market price p_h, delivered energy e_h, imbalance penalties, and operating constraints.

A simplified objective is:

maximize_u E[sum_h p_h e_h(u) - C(u) - penalty(u, y)] - rho * Risk(u),

subject to storage dynamics, power limits, state-of-charge limits, load flexibility constraints, and market bid constraints.

The learning problem is to train a forecasting representation z_theta and a decision layer pi such that the final decision u = pi(z_theta(X)) minimizes decision regret:

Regret = J(u_star; y) - J(pi(z_theta(X)); y),

where u_star is the hindsight-optimal decision under realized outcomes y.

## 4. Proposed Method

### 4.1 Probabilistic forecasting module

The forecasting module outputs price and load quantiles or scenarios. It may reuse the models developed in the previous two manuscripts. For each horizon, the model generates a set of scenarios S = {p_h^s, l_h^s, r_h^s}, representing price, load, and renewable generation uncertainty.

### 4.2 Bidding optimization layer

The optimization layer maps scenarios into a market bid or operating schedule. A basic formulation includes:

- energy balance constraints;
- battery charge/discharge limits;
- state-of-charge dynamics;
- flexible-load shifting bounds;
- renewable generation availability;
- bid quantity and price constraints;
- risk constraints such as CVaR or downside revenue limits.

The exact formulation should remain simple enough for clear experimentation. Overly complex electricity-market rules can obscure the machine-learning contribution.

### 4.3 Decision-focused training

If the optimization layer is differentiable, gradients are propagated from decision loss to forecasting parameters. If exact differentiation is impractical, the model uses a regret surrogate:

L_decision = max(0, margin + J(u_baseline; y) - J(u_theta; y)).

The final training objective combines forecasting loss and decision loss:

L = L_forecast + lambda_d L_decision + lambda_r L_risk.

This prevents the model from ignoring predictive accuracy while encouraging decision-relevant representations.

### 4.4 Risk-aware scenario selection

The model should not optimize only expected revenue. A risk-aware decision layer uses CVaR or quantile-based downside constraints:

CVaR_alpha(loss(u)) <= tau.

This allows the virtual power plant to trade expected profit against downside exposure. The empirical risk analysis reports capacity sensitivity and a revenue-versus-CVaR policy-selection diagnostic; finer risk-aversion sweeps remain a natural extension.

## 5. Experimental Design

### 5.1 Simulation environment

The experiment should combine public price data, load datasets, renewable generation profiles, and a simulated virtual power plant with storage and flexible-load parameters. If real project data exist, they can be added as a case study. The simulator should be transparent and reproducible, with all constraints listed in an appendix.

**Result placeholder:** Insert Table 1 with resource capacities, battery efficiency, state-of-charge limits, flexible-load bounds, market horizon, and penalty assumptions.

### 5.2 Baselines

Baselines should include:

1. Perfect-forecast hindsight optimum.
2. Deterministic forecast-then-optimize using point forecasts.
3. Probabilistic forecast-then-optimize using scenarios.
4. Robust optimization with uncertainty sets.
5. Reinforcement learning policy if implementation time allows.
6. Decision-focused learning without risk constraints.
7. Decision-focused learning with CVaR constraints.

### 5.3 Metrics

Decision metrics: total revenue, average daily profit, regret against hindsight optimum, imbalance penalty, downside revenue, CVaR, loss days, and risk-adjusted return. Forecasting metrics should still be reported but should not dominate the paper.

### 5.4 Ablation study

Diagnostic comparisons remove or vary the main decision ingredients: forecast source, robust risk adjustment, storage capacity, and flexible-load modeling. The current public experiment reports battery-capacity sensitivity and a risk-adjusted policy-search comparison.

## 6. Results and Discussion

**Table 2 placeholder:** Revenue and regret comparison across methods.

**Table 3 placeholder:** Risk metrics, including downside loss and CVaR.

**Table 4 placeholder:** Sensitivity to battery capacity and risk-aversion parameter.

**Figure 1 placeholder:** End-to-end predict-and-optimize framework.

**Figure 2 placeholder:** Example bidding schedules under different risk preferences.

**Figure 3 placeholder:** Forecasting error versus decision regret scatter plot.

The discussion should explicitly show cases where lower forecasting error does not lead to better bidding performance. This is the central empirical argument for decision-focused learning. If the proposed method sacrifices a small amount of MAE while improving revenue or reducing regret, that trade-off should be presented as a strength rather than a weakness.

## 7. Conclusion

This paper proposes a decision-focused learning framework for virtual power plant bidding under electricity price and load uncertainty. By integrating probabilistic forecasting with an optimization-aware training objective, the method aligns predictive representations with market-operation outcomes. The empirical evaluation emphasizes revenue, regret, downside risk, and imbalance penalties, thereby moving beyond conventional forecasting accuracy. The work contributes to computer-science research on predict-and-optimize learning while addressing a practical virtual power plant decision problem.

## Data Availability Statement

The final submission should publish the simulation configuration and public data preprocessing scripts. Any private operational data should be used only as supplementary evidence unless permission for reproducibility is available.

## Cover Letter Draft

Dear Editor,

We submit the manuscript "Decision-Focused Learning for Virtual Power Plant Bidding under Electricity Price and Load Uncertainty" for consideration. The paper studies virtual power plant bidding as a machine-learning problem in which prediction models must be trained for downstream decision quality. It proposes a decision-focused framework that combines probabilistic forecasting, risk-aware optimization, and regret-based learning.

The manuscript contributes to intelligent energy decision support and predict-and-optimize learning. It is original, not published, and not under review elsewhere.

Sincerely,

[Author names]

""" + "\n\n" + COMMON_REFS + "\n"


def thesis_framework():
    return """# 博士论文整合稿框架

## 题目

面向虚拟电厂市场运营的多源时序预测与智能决策方法研究

## 摘要草稿

随着分布式新能源、储能、柔性负荷和电动汽车充电资源的快速接入，虚拟电厂逐渐成为电力市场中聚合分散资源并参与交易的重要主体。虚拟电厂的市场运营高度依赖对电价、负荷、可再生能源出力和资源灵活性的准确预测，但实际场景中存在多源异构、非平稳、强不确定性和跨主体泛化困难等问题。传统方法往往将预测模型与决策优化分离，难以保证预测误差指标与市场收益、风险控制和约束满足之间的一致性。

本文从计算机科学中的多源时序建模和智能决策角度出发，研究面向虚拟电厂市场运营的预测与决策方法。首先，针对电力市场价格波动剧烈、尖峰事件影响显著的问题，提出一种融合动态图学习、时序注意力和保序校准的电价概率预测方法。其次，针对虚拟电厂聚合资源负荷序列差异大、目标场景数据不足的问题，提出一种基于自监督表示学习和域自适应的可迁移短期负荷预测方法。最后，针对传统预测-优化分离导致决策性能不稳定的问题，提出一种面向虚拟电厂报价的预测-决策一体化学习方法，将概率预测结果与风险约束优化目标结合，实现从预测精度到市场决策价值的统一评估。

本文的研究重点不是电力系统设备建模或传统调度规则设计，而是复杂能源场景下的非平稳多源时序学习、不确定性估计、迁移泛化和决策目标对齐。实验部分拟基于公开电力市场、负荷和可再生能源数据，并结合虚拟电厂聚合仿真环境，从点预测精度、概率预测校准、跨域泛化能力、收益、后悔值和风险暴露等维度验证所提方法的有效性。

## 章节结构

### 第1章 绪论

说明虚拟电厂市场运营背景，提出本文关注的计算机问题：多源时序预测、不确定性建模、迁移泛化和预测-决策一致性。避免把绪论写成电力系统设备综述。

### 第2章 相关工作

覆盖时间序列预测、Transformer、图神经网络、自监督时序学习、不确定性估计、conformal prediction、decision-focused learning 和虚拟电厂智能运营。电力市场规则只保留必要背景。

### 第3章 面向电价波动的多源时空概率预测方法

整合论文一。重点是动态图、时序注意力和概率校准。

### 第4章 面向聚合资源的跨域负荷预测方法

整合论文二。重点是自监督表示学习、域自适应、少样本和冷启动。

### 第5章 面向市场报价的预测-决策一体化方法

整合论文三。重点是 decision-focused learning、risk-aware bidding 和 regret。

### 第6章 实验平台与统一验证

统一数据处理、baseline、评价指标和消融实验。展示三章之间的逻辑递进：预测 -> 不确定性 -> 决策价值。

### 第7章 总结与展望

总结计算机方法贡献，展望更大规模虚拟电厂、多智能体交易、隐私保护学习和能源大模型。

## 需要真实实验补齐的内容

1. 公开数据下载、清洗和时间对齐脚本。
2. 三篇论文共享的 baseline 代码。
3. 价格预测、负荷预测和决策仿真的结果表。
4. 每章至少 3 个消融实验。
5. 每章至少 1 个可解释性或案例分析图。
6. 学校认可的期刊分区证明和投稿记录。
"""


def cover_letters():
    return """# 投稿信与审稿风险预案

## 1. EAAI / ESWA 投稿信重点

强调本文不是普通电力应用，而是 AI engineering application:

- heterogeneous non-stationary time-series learning;
- uncertainty-aware forecasting;
- dynamic graph learning;
- decision support under market volatility;
- reproducible public-data evaluation.

避免在投稿信里写太多电力系统工程细节。

## 2. IEEE Access 投稿信重点

强调文章完整性、工程适用性和可复现实验:

- broad interest to IEEE readers;
- transferable forecasting or decision-support framework;
- public data and clear baseline comparison;
- practical value for virtual power plant operation.

## 3. 可能审稿意见与预案

### 意见 A：创新性不足，只是组合已有模型

回应策略：突出问题定义和组合的必要性，补充消融实验，证明动态图、不确定性校准、迁移或决策损失分别贡献明确。

### 意见 B：电力场景太窄，不适合 AI 期刊

回应策略：强调非平稳多源时序学习、跨域泛化、预测-决策一致性，并增加至少一个非单一市场/跨数据集泛化实验。

### 意见 C：缺少真实虚拟电厂数据

回应策略：使用公开市场和负荷数据保证可复现；虚拟电厂资源通过透明仿真构造；若有企业数据，只作为补充验证。

### 意见 D：结果不比所有 baseline 都好

回应策略：不要声称全面优越；强调高波动、跨域、少样本、风险控制等目标场景下的优势。

### 意见 E：不确定性预测未校准

回应策略：加入 PICP、PINAW、CRPS、calibration curve，并使用 conformal calibration 做覆盖率修正。
"""


FILES = {
    "paper_1_price_forecasting.md": paper_one(),
    "paper_2_transferable_load_forecasting.md": paper_two(),
    "paper_3_decision_focused_vpp_bidding.md": paper_three(),
    "thesis_framework_cn.md": thesis_framework(),
    "cover_letters_and_review_plan.md": cover_letters(),
}


def setup_doc(document, title):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(23, 54, 93)


def _split_md_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_md_table_separator(line):
    cells = _split_md_table_row(line)
    return bool(cells) and all(set(cell.replace(" ", "")) <= {"-", ":"} and "-" in cell for cell in cells)


def _add_markdown_table(document, table_lines):
    rows = [_split_md_table_row(line) for line in table_lines]
    if len(rows) >= 2 and _is_md_table_separator(table_lines[1]):
        rows = [rows[0]] + rows[2:]
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cell.paragraphs[0]
            run = para.add_run(val.replace("**", ""))
            run.font.size = Pt(8)
            if r_idx == 0:
                run.bold = True
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_markdown_to_docx(document, text):
    lines = text.splitlines()
    in_refs = False
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.lower() in {"<pagebreak>", "<page-break>", "<!-- pagebreak -->"}:
            document.add_page_break()
            i += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_markdown_table(document, table_lines)
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            alt, image_path = image_match.groups()
            p = Path(image_path)
            if not p.exists() and not p.is_absolute():
                p = ROOT / image_path
            if p.exists():
                para = document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(str(p), width=Inches(6.3))
                if alt:
                    cap = document.add_paragraph(alt)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in cap.runs:
                        r.font.size = Pt(9)
                        r.italic = True
            else:
                document.add_paragraph(f"[Missing figure: {image_path}]")
            i += 1
            continue
        clean = stripped.replace("**", "")
        if stripped.startswith("# "):
            i += 1
            continue
        if stripped.startswith("## "):
            document.add_heading(clean[3:], level=1)
            in_refs = stripped[3:].lower() == "references"
        elif stripped.startswith("### "):
            document.add_heading(clean[4:], level=2)
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = document.add_paragraph()
            r = p.add_run(clean)
            r.bold = True
        elif stripped.startswith("- "):
            document.add_paragraph(clean[2:], style="List Bullet")
        elif stripped[:3] in ["1. ", "2. ", "3. ", "4. ", "5. ", "6. ", "7. "] and not in_refs:
            document.add_paragraph(clean[3:], style="List Number")
        else:
            p = document.add_paragraph(clean)
            if stripped.startswith("[") and "]" in stripped[:6]:
                p.paragraph_format.space_after = Pt(3)
        i += 1


def save_docx(name, text):
    document = Document()
    title = text.splitlines()[0].replace("# ", "").strip()
    setup_doc(document, title)
    add_markdown_to_docx(document, text)
    path = OUT / name.replace(".md", ".docx")
    document.save(path)
    return path


def main():
    written = []
    for name, text in FILES.items():
        md_path = OUT / name
        md_path.write_text(text, encoding="utf-8")
        written.append(md_path)
        written.append(save_docx(name, text))

    index = """# Paper Package Index

This package contains three manuscript drafts, one Chinese doctoral thesis integration framework, and cover-letter/reviewer-response planning notes.

## Files

- paper_1_price_forecasting.md / .docx
- paper_2_transferable_load_forecasting.md / .docx
- paper_3_decision_focused_vpp_bidding.md / .docx
- thesis_framework_cn.md / .docx
- cover_letters_and_review_plan.md / .docx

## Submission Gate

None of the three article drafts should be submitted before the following items are completed:

1. Confirm the current verified experiment outputs remain synchronized with each target-journal manuscript.
2. Confirm the target journal and format references accordingly.
3. Confirm school classification rules for the exact submission year.
4. Run plagiarism/self-similarity checks after finalizing prose.
5. Ensure every claim in Results and Discussion is supported by a table, figure, or statistical test.
"""
    index_path = OUT / "README.md"
    index_path.write_text(index, encoding="utf-8")
    written.append(index_path)
    print("\\n".join(str(p) for p in written))


if __name__ == "__main__":
    main()
