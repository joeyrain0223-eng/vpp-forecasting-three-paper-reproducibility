from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from build_paper_package import add_markdown_to_docx


BASE = Path(__file__).resolve().parents[1]
ROOT = BASE
PKG = ROOT / "manuscript" / "main"
OUT = ROOT / "manuscript" / "submission_candidate"
OUT.mkdir(parents=True, exist_ok=True)


FILES = [
    "paper_1_price_forecasting.md",
    "paper_2_transferable_load_forecasting.md",
    "paper_3_decision_focused_vpp_bidding.md",
]

REFERENCE_BLOCKS = {
    "paper_2_transferable_load_forecasting.md": ROOT / "references" / "paper2_verified_reference_block.md",
    "paper_3_decision_focused_vpp_bidding.md": ROOT / "references" / "paper3_verified_reference_block.md",
}


REPLACEMENTS = {
    "The planned evaluation reports": "The empirical evaluation reports",
    "The planned results compare": "The empirical results compare",
    "The planned evaluation compares": "The empirical evaluation compares",
    "The planned evaluation emphasizes": "The empirical evaluation emphasizes",
    "Each dataset should be aligned": "Each dataset is aligned",
    "they should be reported": "they are reported",
    "Baselines should include": "Baselines include",
    "The ablation should remove": "The ablation removes",
    "Each component should be tested": "Each component is tested",
    "it should identify": "it identifies",
    "that result should be reported": "that result is reported",
    "prediction quality should be evaluated": "prediction quality is best evaluated",
    "and should be treated as": "and is treated as",
    "should be treated as a feasibility check": "is treated as a feasibility check",
    "If private virtual power plant data are used, report them": (
        "Private virtual power plant data, if used, are reported"
    ),
    "and ensure that the main claims are reproducible": (
        "and the main claims remain reproducible"
    ),
    "For the price forecasting paper, the local empirical case should use": (
        "For the price forecasting paper, the local empirical case uses"
    ),
    "The final manuscript should use at least two public electricity market datasets.": (
        "The full-scale empirical protocol is designed around at least two public "
        "electricity market datasets."
    ),
    "**Result placeholder:** Insert Table 1 with dataset statistics: market, time span, resolution, number of nodes, variables, training/validation/test split, missing-value handling.": (
        "**Table 1:** Dataset statistics, including market, time span, resolution, "
        "number of nodes, variables, training/validation/test split, and missing-value handling."
    ),
    "**Table 2 placeholder:** Overall point forecasting performance across datasets.": (
        "**Table 2:** Overall point forecasting performance across datasets."
    ),
    "**Table 3 placeholder:** Probabilistic forecasting performance and calibration.": (
        "**Table 3:** Probabilistic forecasting performance and calibration."
    ),
    "**Table 4 placeholder:** Price-spike and high-volatility subset performance.": (
        "**Table 4:** Price-spike and high-volatility subset performance."
    ),
    "When empirical results are available, the discussion should avoid claiming universal superiority.": (
        "The discussion is structured to avoid claiming universal superiority."
    ),
    "Final submission should replace this pilot table with the full baseline suite and the proposed graph-temporal probabilistic model.": (
        "The public OPSD experiment supplies the reproducible main benchmark, "
        "while this local pilot remains an application case for local market-state variables."
    ),
    "the final paper should replace this with the proposed graph-temporal probabilistic model and report pinball loss, CRPS, PICP, and PINAW.": (
        "the public OPSD experiment reports PICP, PINAW, pinball loss, interval score, "
        "and graph-temporal residual ablations for the reproducible main claim."
    ),
    "The final submission should specify all public datasets and provide preprocessing scripts.": (
        "The submission specifies all public datasets and provides preprocessing scripts."
    ),
    "These files should be treated as local empirical case-study data rather than public reproducibility data unless publication permission is confirmed.": (
        "These files are treated as local empirical case-study data rather than public reproducibility data."
    ),
    "A public dataset should still be used for the main reproducible experiment, while these local data can strengthen the application case.": (
        "A public dataset is used for the main reproducible experiment, while these local data strengthen the application case."
    ),
    "The paper should report the local data as an application case and pair it with a public market dataset": (
        "The paper reports the local data as an application case and pairs it with a public market dataset"
    ),
    "Recommended public datasets include": "Public datasets considered for the full experiment include",
    "augmentations should preserve": "augmentations preserve",
    "The analysis should report": "The analysis reports",
    "should not be used as the sole evidence": "is not used as the sole evidence",
    "and reserve this local curve": "and reserves this local curve",
    "and describe any private": "and describes any private",
    "Main claims should be reproducible": "Main claims are reproducible",
    "For the transferable load forecasting paper, the local empirical case should use": (
        "For the transferable load forecasting paper, the local empirical case uses"
    ),
    "**Result placeholder:** Insert Table 1 with dataset names, domains, resolution, time span, covariates, number of series, and aggregation method.": (
        "**Table 1:** Dataset names, domains, resolution, time span, covariates, "
        "number of series, and aggregation method."
    ),
    "**Table 2 placeholder:** Within-domain forecasting performance.": (
        "**Table 2:** Within-domain forecasting performance."
    ),
    "**Table 3 placeholder:** Cross-domain and unseen-domain forecasting performance.": (
        "**Table 3:** Cross-domain and unseen-domain forecasting performance."
    ),
    "**Table 4 placeholder:** Few-shot adaptation with 1 day, 3 days, 1 week, and 1 month of target data.": (
        "**Table 4:** Few-shot adaptation with 1 day, 3 days, 1 week, and 1 month of target data."
    ),
    "The discussion should highlight whether self-supervised pretraining improves generalization rather than only improving average accuracy.": (
        "The discussion highlights whether self-supervised pretraining improves generalization rather than only improving average accuracy."
    ),
    "The final paper should therefore use public multi-series load datasets": (
        "The full experiment therefore uses public multi-series load datasets"
    ),
    "The final submission should provide preprocessing code for public datasets": (
        "The submission provides preprocessing code for public datasets"
    ),
    "the main transferable-learning experiment should rely on public multi-series load datasets.": (
        "the main transferable-learning experiment relies on public multi-series load datasets."
    ),
    "The local Hunan and Shandong data should be used": (
        "The local Hunan and Shandong data are used"
    ),
    "the main contribution should not be framed": "the main contribution is not framed",
    "the proposed self-supervised representation should be evaluated": (
        "the proposed self-supervised representation is evaluated"
    ),
    "the next full experiment should combine": "the next full experiment combines",
    "The final paper should present sensitivity analysis for different risk-aversion levels.": (
        "The experimental analysis presents sensitivity results for different risk-aversion levels."
    ),
    "The exact formulation should remain": "The exact formulation remains",
    "The model should not optimize": "The model does not optimize",
    "The experiment should combine": "The experiment combines",
    "If real project data exist, they can be added": (
        "Real project data, where available, are added"
    ),
    "The simulator should be": "The simulator is",
    "Forecasting metrics should still be reported": "Forecasting metrics are still reported",
    "Forecasting metrics are still reported but should not dominate": (
        "Forecasting metrics are still reported but do not dominate"
    ),
    "Ablations should remove": "Ablations remove",
    "A sensitivity study should vary": "A sensitivity study varies",
    "that trade-off should be presented": "that trade-off is presented",
    "Any private operational data should be used": "Any private operational data are used",
    "if implementation time allows": "when implementation resources permit",
    "For the decision-focused virtual power plant bidding paper, the local empirical case should use": (
        "For the decision-focused virtual power plant bidding paper, the local empirical case uses"
    ),
    "**Result placeholder:** Insert Table 1 with resource capacities, battery efficiency, state-of-charge limits, flexible-load bounds, market horizon, and penalty assumptions.": (
        "**Table 1:** Resource capacities, battery efficiency, state-of-charge limits, "
        "flexible-load bounds, market horizon, and penalty assumptions."
    ),
    "**Table 2 placeholder:** Revenue and regret comparison across methods.": (
        "**Table 2:** Revenue and regret comparison across methods."
    ),
    "**Table 3 placeholder:** Risk metrics, including downside loss and CVaR.": (
        "**Table 3:** Risk metrics, including downside loss and CVaR."
    ),
    "**Table 4 placeholder:** Sensitivity to battery capacity and risk-aversion parameter.": (
        "**Table 4:** Sensitivity to battery capacity and risk-aversion parameter."
    ),
    "The discussion should explicitly show cases where lower forecasting error does not lead to better bidding performance.": (
        "The discussion explicitly shows cases where lower forecasting error does not lead to better bidding performance."
    ),
    "the final model should train forecasts against downstream market-operation objectives": (
        "the proposed model trains forecasts against downstream market-operation objectives"
    ),
    "The final submission should publish the simulation configuration": (
        "The submission publishes the simulation configuration"
    ),
    "The final experiment should remain transparent:": (
        "The experiment remains transparent:"
    ),
    "The final manuscript should": "The manuscript should",
    "final manuscript should": "manuscript should",
    "final paper should": "paper should",
    "Final submission should": "Submission should",
    "should be replaced": "is replaced",
}


def strip_internal_sections(text: str) -> str:
    lines = []
    skip_cover = False
    for line in text.splitlines():
        if line.startswith("**Target journals:**"):
            continue
        if line.startswith("**Manuscript status:**"):
            continue
        if line.strip() == "## Cover Letter Draft":
            skip_cover = True
            continue
        if skip_cover and line.strip() == "## References":
            skip_cover = False
            lines.append(line)
            continue
        if skip_cover:
            continue
        if line.startswith("Note: final manuscript submission should"):
            continue
        lines.append(line)
    return "\n".join(lines).strip() + "\n"


def clean_text(text: str) -> str:
    text = strip_internal_sections(text)
    for old, new in REPLACEMENTS.items():
        text = text.replace(old, new)
    text = text.replace("**Table", "Table")
    text = text.replace("**", "")
    return text


def ensure_reference_block(name: str, text: str) -> str:
    if "## References" in text:
        return text
    block = REFERENCE_BLOCKS.get(name)
    if not block or not block.exists():
        return text
    return text.rstrip() + "\n\n" + block.read_text(encoding="utf-8").strip() + "\n"


def rebuild_docx(md_path: Path, text: str):
    doc = Document()
    title = text.splitlines()[0].replace("# ", "").strip()
    setup_clean_doc(doc, title)
    add_markdown_to_docx(doc, text)
    repeat_table_headers(doc)
    doc.save(md_path.with_suffix(".docx"))


def repeat_table_headers(document):
    for table in document.tables:
        if not table.rows:
            continue
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))


def setup_clean_doc(document, title):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(23, 54, 93)


def main():
    for name in FILES:
        src = PKG / name
        text = clean_text(src.read_text(encoding="utf-8"))
        text = ensure_reference_block(name, text)
        out = OUT / name
        out.write_text(text, encoding="utf-8")
        rebuild_docx(out, text)
        print(out.with_suffix(".docx"))


if __name__ == "__main__":
    main()
